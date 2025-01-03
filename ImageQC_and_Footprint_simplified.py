import docx
from docx.shared import Pt
import os 
import arcpy
import xml.etree.ElementTree as ET
import fnmatch
import rasterio
import numpy as np
import geopandas as gpd
import os 
import re
from utils import cems_utils
from osgeo import gdal, osr
from shapely.geometry import Polygon, MultiPolygon, mapping
from functools import reduce
from osgeo import ogr
import tkinter as tk
from tkinter import ttk
from datetime import datetime

class MetadataReader:

    def __init__(self, sensor_folder):
        self.sensor_folder = sensor_folder.split(";")[0]
    
    def _parse_xml(self, xml_file):
        tree = ET.parse(xml_file)
        root = tree.getroot()
        return root
    
    def _format_datetime(self, date_time):
        if not date_time:
            arcpy.AddError("Date and Time not found in metadata.")
        return f"{date_time[8:10]}/{date_time[5:7]}/{date_time[0:4]} {date_time[11:13]}:{date_time[14:16]}"
    
    def Sentinel2(self):
        Date_and_Time = ""
        for root, subdirs, files in os.walk(self.sensor_folder):
            for d in files:
                if d.endswith('2A.xml'):
                    xmlfile = os.path.join(root,d)
        root = self._parse_xml(xmlfile)
        for e in root:
             for attr in e:
                if attr.tag == 'Product_Info':
                    for l in attr:
                        if l.tag == 'PRODUCT_START_TIME':
                            Date_and_Time = l.text
        DateTime = self._format_datetime(Date_and_Time)
        return (DateTime)
    
    def Sentinel1(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for d in files:
                if d.endswith('manifest.safe'):
                    xmlfile = os.path.join(root,d)
        root = self._parse_xml(xmlfile)
        for b in root:
            if b.tag == 'metadataSection':
                for c in b:
                    for d in c:
                        for e in d:
                            for f in e:
                                for g in f:
                                    if 'stopTime' in g.tag:
                                        date_time = g.text
        DateTime = self._format_datetime(date_time)
        return(DateTime)
    
    def PAZ_and_Planetscope_and_SkySat_and_TERRASARX(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                if fnmatch.fnmatch(file_name, "INSPIRE.xml"):
                    xml_file = os.path.join(root, file_name)
        root = self._parse_xml(xml_file)
        end_position_elem = root.find('.//{http://www.opengis.net/gml/3.2}endPosition')
        acquisition_time = end_position_elem.text
        DateTime = self._format_datetime(acquisition_time)
        return(DateTime)
    
    def Pleiads_neo(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                if fnmatch.fnmatch(file_name, "VOL_PNEO.XML"):
                    xml_file = os.path.join(root, file_name)               
        root = self._parse_xml(xml_file)
        product_information = root.find(".//Product_Information")
        delivery_identification = product_information.find(".//Delivery_Identification")
        acquisition_time = delivery_identification[0].text
        DateTime = self._format_datetime(acquisition_time)
        return(DateTime)
    
    def WorldView_and_Geoeye(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                if file_name.endswith("README.XML"):
                    xml_file = os.path.join(root, file_name)       
        root = self._parse_xml(xml_file)
        acquisition_time = root.find(".//COLLECTIONSTOP").text
        DateTime = self._format_datetime(acquisition_time)
        return(DateTime)
    
    def CosmoSkyMed(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                if file_name.endswith("AccompanyingSheet.xml"):
                    xml_file = os.path.join(root, file_name)
        root = self._parse_xml(xml_file)
        acquisition_time = root.find(".//SensingStopTime").text
        DateTime = self._format_datetime(acquisition_time)
        return(DateTime)
    
    def SPOT_and_Pleiads(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                if file_name.startswith("DIM"):
                    xml_file = os.path.join(root, file_name)
        root = self._parse_xml(xml_file)
        acquisition_time = root.find(".//IMAGING_TIME").text
        acquisition_date = root.find(".//IMAGING_DATE").text
        acquisition_string = acquisition_date + "T" + acquisition_time
        DateTime = self._format_datetime(acquisition_string)
        return (DateTime)
    
    def Landsat8_9(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                    if file_name.endswith(".xml"):
                        xml_file = os.path.join(root, file_name)               
        root = self._parse_xml(xml_file)
        acquisition_time = root.find(".//DATE_PRODUCT_GENERATED").text
        DateTime = self._format_datetime(acquisition_time)
        return(DateTime)
    
    def RadarSat2(self):
        for root, subdirs, files in os.walk(self.sensor_folder):
            for file_name in files:
                if  fnmatch.fnmatch(file_name, "product.xml"):
                    xml_file = os.path.join(root, file_name)
        root = self._parse_xml(xml_file)
        for b in root:
            if 'sourceAttributes' in b.tag:
                for c in b:
                    if 'rawDataStartTime' in c.tag:
                        acquisition_time = c.text
        DateTime = self._format_datetime(acquisition_time)
        return(DateTime)
    
class QualityCheckCompiler:

    def __init__(self, EMSR_Number, AOI, Map_Type, File_name, DateTime, Availability, 
        Reception, res, output_path, productionSite):

        self.EMSR_Number = EMSR_Number
        self.AOI = AOI
        self.Map_Type = Map_Type
        self.File_name = File_name
        self.DateTime = DateTime
        self.Availability = Availability
        self.Reception = Reception
        self.res = res
        self.output_path = output_path
        self.productionSite = productionSite
    
    def Compiler(self):
        doc_input = os.path.join(os.path.dirname(os.path.realpath(__file__)), "utils\\EMSRxxx_AOIxx_MapType_PRODUCT_sensorName_yyyymmdd_hhmm_IMAGE_QC_v1.docx")
        doc = docx.Document(doc_input)
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(12)
        font.name = 'Arial'
        para = doc.paragraphs
        Production_par = para[6].text
        EMSR_par = para[7].text
        AOI_par = para[8].text
        MapType_par = para[9].text
        Identifier_par = para[11].text
        Acquisition_par = para[14].text
        Avaiability_par = para[15].text
        Reception_par = para[16].text
        Acceptance_par = para[17].text
        Res_par = para[18].text
        lista2 = [Production_par, EMSR_par,AOI_par,MapType_par,Identifier_par,Acquisition_par,Avaiability_par,Reception_par,Acceptance_par, Res_par]
        lista = ['Production', 'EMSR','AOI','Map_Type','File_Name','Acquisition','Availability','Reception','Acceptance','Res']
        x = 0
        Dictionary_QC = {
        'Production': 'Production site (PS): ' + str(self.productionSite),
        'EMSR': f"Activation number: {self.EMSR_Number}",
        'AOI':'AOI N°: ' + str(self.AOI),
        'Map_Type':'Map Type: '+ str(self.Map_Type),
        'File_Name':'File identifier: ' + str(self.File_name) +'.tif',
        'Acquisition': 'Acquisition Date and Time: ' + str(self.DateTime) + ' UTC',
        'Availability': 'Availability Date and Time: ' + str(self.Availability)  + ' UTC',
        'Reception': 'Reception Date and Time: ' + str(self.Reception) + ' UTC',
        'Acceptance': 'Image Quality Acceptance Date and Time: ' + datetime.utcnow().strftime("%d/%m/%Y %H:%M") + ' UTC',
        'Res': 'Resolution: ' + str(self.res) + " m"
        }
        for p in doc.paragraphs:
            paragraphs = doc.paragraphs
            try:
                if lista2[x] in p.text:
                    for i in range(len(paragraphs)):
                        if lista2[x] in paragraphs[i].text:      
                            text = paragraphs[i].text.replace(lista2[x], Dictionary_QC[lista[x]])
                            paragraphs[i].text = text
                            x = x+1
            except:
                pass
        doc_out = reduce(os.path.join, [self.output_path, self.File_name + '_IMAGE_QC.docx'])
        arcpy.AddMessage(f"Filling Quality Check .docx file at: {doc_out}")
        doc.save(doc_out)

class ImageFootprint:

    def __init__(self, image_path, raw_sensor_folder, preOrthoFolder, postOrthoFolder, shapefile_footprint_path, 
        kml_footprint_path ,A2_image_footprint_a, A0_source_table, gdb_path,
        pre_image_condition, sensor_user_input, DateTime, res, 
        erasing_condition, mosaick_name, sensor_domain_value, SAR, activation_crs):
       
       self.image_path = image_path
       self.raw_sensor_folder = raw_sensor_folder
       self.shp_footprint_path = shapefile_footprint_path
       self.A2_image_footprint_a = A2_image_footprint_a 
       self.A0_source_table = A0_source_table
       self.gdb_path = gdb_path
       self.pre_image_condition = pre_image_condition
       self.sensor_user_input = sensor_user_input
       self.Date = DateTime[:10]
       self.Time = f"T{DateTime[11:]}:00Z"
       self.res = res
       self.erasing_condition = erasing_condition
       self.mosaick_name = mosaick_name
       self.preOrthoFolder = preOrthoFolder
       self.postOrthoFolder = postOrthoFolder
       self.SAR = SAR
       self.activation_crs = f"EPSG:{activation_crs}"
       self.mosaick_path = (os.path.join(self.preOrthoFolder, f"{self.mosaick_name}.tif") if self.pre_image_condition 
       else os.path.join(self.postOrthoFolder, f"{self.mosaick_name}.tif"))
       self.kml_footprint_path = kml_footprint_path
       self.sensor_domain_value = sensor_domain_value
       self.temp_folder = cems_utils.createTempFolder()
    
    def A0SourceCompiler(self):
        attribute1 = "eventphase"
        attribute2 = "src_id" 
        eventphase = 1 if self.pre_image_condition == True else 2
        index = 0
        src_id = 1
        tuples_list = []
        with arcpy.da.UpdateCursor(self.A0_source_table, ["OID@", attribute1, attribute2]) as cursor:
            for row in cursor:
                if (row[1] == 1 and row[2] < 10):
                    tuples_list.append((row[2], False))
                elif (row[1] == 2):
                    tuples_list.append((row[2], True))    
                    if self.erasing_condition == True:
                        arcpy.AddMessage("Deleting previous record from A0 source...")
                        cursor.deleteRow()
                index += 1
        if tuples_list:
            src_id = tuples_list[-1][0] + 1
        fields = ["OBJECTID", "source_nam", "src_date", "source_tm", "sensor_res", "eventphase", "src_id"]
        new_row_values = [index, self.sensor_domain_value[self.sensor_user_input], self.Date, self.Time, f"{self.res} m", eventphase, src_id]
        with arcpy.da.InsertCursor(self.A0_source_table, fields) as cursor:
            arcpy.AddMessage("Adding new record to A0 source...")
            for i in range(index + 1):
                if i == index:
                    cursor.insertRow(new_row_values)
        return(src_id, tuples_list)
        
    def A2_eraser(self, tuples_list):
        with arcpy.da.UpdateCursor(self.A2_image_footprint_a, ["OID@", "or_src_id"]) as cursor:
            for row in cursor:
                for src_id, to_delete in tuples_list:
                    if row[1] == src_id and to_delete == True:
                        if self.erasing_condition:
                            arcpy.AddMessage("Deleting previous post-Image Footprints from A2 layer...")
                            cursor.deleteRow()

    def Composite(self):
        if self.sensor_user_input in ["Sentinel-2", "Landsat-8", "Landsat-9"]:
            separate_folders = self.raw_sensor_folder.split(";")
            extention = (("B02_10m.jp2", "B03_10m.jp2", "B04_10m.jp2", "B08_10m.jp2", "B12_20m.jp2") if self.sensor_user_input == "Sentinel-2"
            else ("B2.TIF", "B3.TIF", "B4.TIF", "B5.TIF", "B6.TIF"))
            composite_options = {"VRTNodata": 0, "separate": True, "resolution": "highest"}
            i = 1 
            VRTs = []
            for folder in separate_folders:
                arcpy.AddMessage(f"Making composite bands of folder number {i}...")
                tiles = []
                for root, subdirs, files in os.walk(folder):
                    for file in files:
                        if file.endswith(extention):
                            file_path = os.path.join(root, file)
                            tiles.append(file_path)
                try:
                    tiles[0], tiles[2] = tiles[2], tiles[0] # To make the composite as JRC requests
                except:
                    arcpy.AddError(f"Your {self.sensor_user_input} folder contains too many characters and the tool could not find the tiles. Please shorten some paths by renaming the folder or by reducing the levels inside the folder!")
                storing_vrt_path = os.path.join(self.temp_folder, f"Composite_tile{i}.vrt")
                gdal.BuildVRT(storing_vrt_path, tiles, **composite_options)
                VRTs.append(storing_vrt_path)
                arcpy.AddMessage(f"Composite of folder number {i} done!")
                i += 1
            return(VRTs)
        else: return(None)
    
    def GetImageCrs(self, image):
        raster = gdal.Open(image)
        crs_wkt = raster.GetProjection()
        spatial_ref = osr.SpatialReference()
        spatial_ref.ImportFromWkt(crs_wkt)
        spatial_ref.AutoIdentifyEPSG()
        epsg_code = spatial_ref.GetAuthorityCode('PROJCS')
        if epsg_code is None:
            epsg_code = spatial_ref.GetAuthorityCode('GEOGCS')
        image_crs = f"EPSG:{epsg_code}"
        return image_crs

    def Reproject(self, n_images):
        images_list = []
        i = 1
        for image in n_images:
            image_crs = self.GetImageCrs(image)
            if image_crs != self.activation_crs:
                arcpy.AddMessage(f"Tile {i} crs and activation crs do not coincide! Proceeding with reprojection...")
                storing_vrt_path = os.path.join(self.temp_folder, f"Reprojected_tile{i}.vrt")
                option = {"format" : "VRT", "dstSRS": self.activation_crs}
                reprojection_result = gdal.Warp(storing_vrt_path ,image, **option)
                if reprojection_result is None:
                    arcpy.AddError(f"The automatic reprojection of tile {i} could not reproject the image crs from {image_crs} to {self.activation_crs}. Please reproject manually and re-run the tool!")
                images_list.append(storing_vrt_path)
                arcpy.AddMessage(f"Reprojection of tile {i} successful!")
            else:
                images_list.append(image)
            i += 1
        return images_list
                
    def MosaickAndExport(self, composite_images):
        if self.sensor_user_input in ["Sentinel-2", "Landsat-8", "Landsat-9"]:
            n_images = composite_images
        else:
            n_images = self.image_path.split(";")
        if len(n_images) > 1:
            arcpy.AddMessage("Mosaicking images...")
            mosaicking_images = self.Reproject(n_images)
            option = {"srcNodata": 0,"VRTNodata": 0, "resolution": "highest"}
            vrt = os.path.join(self.temp_folder , self.mosaick_name) + ".vrt"
            gdal.BuildVRT(vrt, mosaicking_images, **option)
            option2 = {"format" : "GTiff"}    
            gdal.Translate(self.mosaick_path, vrt, **option2)
            arcpy.AddMessage(f"Mosaick of {self.sensor_user_input} done! Saved at: {self.mosaick_path}")
        elif len(n_images) == 1 and self.sensor_user_input in ["Sentinel-2", "Landsat-8", "Landsat-9"]:
            option = {"format" : "GTiff"}
            storing_vrt_path = os.path.join(self.temp_folder, "Composite_tile1.vrt")
            gdal.Translate(self.mosaick_path, storing_vrt_path, **option)
            arcpy.AddMessage(f"{self.sensor_user_input} image saved at {self.mosaick_path}")
        elif len(n_images) == 1 and n_images[0].endswith((".TIL", ".xml",".JP2")):
            translating_image = self.Reproject(n_images)
            option = {"format" : "GTiff"}
            arcpy.AddMessage(f"Transalting tile: {translating_image}")
            try:
                gdal.Translate(self.mosaick_path ,translating_image[0], **option)
            except Exception as e:
                arcpy.AddError(e)
        elif self.SAR or (len(n_images) == 1 and n_images[0].endswith(".tif")):
            projected_image = self.Reproject(n_images)
            if projected_image[0] == n_images[0]:
                arcpy.AddMessage(f"Your image is currently located at {self.image_path}. Make sure this is the right path!")
                self.mosaick_path = self.image_path
            else:
                option = {"format" : "GTiff"}
                gdal.Translate(self.mosaick_path, projected_image[0], **option)
                arcpy.AddMessage(f"Your image had to be reprojected and it has been located at {self.mosaick_path}")

    def RemoveInvalidData(self):
        if self.SAR:
            arcpy.AddMessage("Ckecking for invalid data presence in your SAR image...")
            raster = rasterio.open(self.mosaick_path)
            band = raster.read(1)
            if np.isnan(np.sum(band)):
                arcpy.AddMessage("Your image has invalid data! Handling them...")
                metadata = raster.meta
                metadata["nodata"] = 0.0 
                processed_band = np.nan_to_num(band, nan = 0.0)
                processed_raster_path = os.path.join(os.path.dirname(self.shp_footprint_path), self.mosaick_name + "_modified.tif")
                with rasterio.open(processed_raster_path, 'w', **metadata) as dst:
                    dst.write(processed_band.astype(rasterio.float32), 1)
                arcpy.AddMessage(f"Invalid data have been handled. Find processed SAR image at: {processed_raster_path}")
                self.mosaick_path = processed_raster_path
            else:
                arcpy.AddMessage("Your image does not have invalid data. Proceeding with next operations...")

    def ShapefileFootprint(self):
        option = {'srcNodata': float(0)}
        arcpy.AddMessage("Creating image footprint...")
        gdal.Footprint(self.shp_footprint_path, self.mosaick_path, format='ESRI Shapefile', **option)
        # Fill Potential holes:
        shapefile = gpd.read_file(self.shp_footprint_path)
        arcpy.AddMessage("Filling eventual holes inside image footprint...")
        geometries = shapefile['geometry']
        no_holes = [MultiPolygon([Polygon(geom.exterior)]) if geom.geom_type == 'Polygon'
        else MultiPolygon([Polygon(p.exterior) for p in geom.geoms])
        for geom in geometries]
        shapefile['geometry'] = no_holes
        # Explode to get Single-part polygons
        arcpy.AddMessage("Converting Multipolygon into Polygon...")
        exploded = shapefile.explode(column = 'geometry')
        exploded['area'] = exploded.area
        max_area = exploded['area'].max()
        footprint = exploded[exploded['area'] == max_area]
        footprint = footprint.drop(columns = ['area'])
        footprint['obj_type'] = 1  
        footprint.to_file(self.shp_footprint_path, driver='ESRI Shapefile')
        arcpy.AddMessage(f"Shapefile footprint successfully generated at {self.shp_footprint_path}")
    
    def UploadOnToc(self):

        def press_yes_button():
            nonlocal abort_procedure
            root.destroy()
            aprx.deleteItem(map)
            abort_procedure = False
            
        def press_no_button():
            nonlocal abort_procedure
            aprx.deleteItem(map)
            root.destroy()
            abort_procedure = True

        aprx = arcpy.mp.ArcGISProject("current")
        aprx.createMap("Image QC Map", 'MAP')
        map = aprx.listMaps("Image QC Map")[0]
        map.addDataFromPath(self.mosaick_path)
        map.addDataFromPath(self.shp_footprint_path)
        aoi_layer = map.listLayers(self.shp_footprint_path.split('\\')[-1].replace('.shp', ''))[0]
        aoi_layer.definitionQuery = 'obj_type = 1'
        aoi_symbology = aoi_layer.symbology
        aoi_symbology.renderer.symbol.applySymbolFromGallery("Black Outline")
        aoi_symbology.renderer.symbol.outlineColor = {'RGB': [0, 255, 0, 255]}
        aoi_symbology.renderer.symbol.outlineWidth = 2
        aoi_layer.symbology = aoi_symbology
        map.openView()
        abort_procedure = None
        root = tk.Tk()
        root.title("Accept ImageFootprint")
        root.configure(bg="lightblue")
        root.update_idletasks()
        label = tk.Label(root, text = "Do you accept the image and the relative footprint?", font=("Helvetica", 14), bg="lightblue")
        label.pack(pady=10)
        button_frame = tk.Frame(root, bg="lightblue")
        button_frame.pack(pady=10)
        yes_button = tk.Button(button_frame, text="Yes", command=press_yes_button, font=("Helvetica", 12), bg="green", fg="white", padx=20, pady=10)
        yes_button.pack(side=tk.LEFT, padx=20)
        no_button = tk.Button(button_frame, text="No", command=press_no_button, font=("Helvetica", 12), bg="red", fg="white", padx=20, pady=10)
        no_button.pack(side=tk.RIGHT, padx=20)
        root.mainloop()
        return abort_procedure
    
    def FindNoData(self, or_src_id):
        footprint = gpd.read_file(self.shp_footprint_path)
        footprint = footprint.to_crs(4326)  
        footprint['or_src_id'] = int(or_src_id)
        aoi = gpd.read_file(self.gdb_path, layer = "A1_area_of_interest_a")
        arcpy.AddMessage("Checking if Image Footprint covers the entire AOI...")
        erase = aoi.overlay(footprint, how = "difference")
        if len(erase) > 0:
            arcpy.AddMessage("NoData areas found!")
            footprint.at[1, 'geometry'] = erase.at[0, 'geometry']
            footprint["obj_type"][1] = 3
            footprint["or_src_id"][1] = int(or_src_id)
        else:
            arcpy.AddMessage("Image Footprint covers all the Area Of Interest")
        footprint.to_file(self.shp_footprint_path)
    
    def GDBFootprint(self):
        temp_gdb = cems_utils.createTempGdb()
        arcpy.conversion.FeatureClassToGeodatabase(
        Input_Features= self.shp_footprint_path,
        Output_Geodatabase= temp_gdb
        )
        feature_class_path = os.path.join(temp_gdb, self.shp_footprint_path.split('\\')[-1].replace('.shp', ''))
        arcpy.AddMessage("Appending data to A2_Image_Footprint_a feature class...")
        arcpy.management.Append(
        inputs= feature_class_path,
        target=self.A2_image_footprint_a,
        schema_type="NO_TEST",
        field_mapping=r'obj_type "Object Type" true true false 4 Long 0 0,First,#,feature_class_path,obj_type,-1,-1;or_src_id "Origin Source Identifier" true true false 4 Long 0 0,First,#,feature_class_path,or_src_id,-1,-1',
        subtype="",
        expression="",
        match_fields=None,
        update_geometry="NOT_UPDATE_GEOMETRY"
        )
    
    def KmlFootprint(self):
        arcpy.AddMessage("Generating .kml Image Footrpint...")
        driver_shp = ogr.GetDriverByName('ESRI Shapefile')
        dataset_shp = driver_shp.Open(self.shp_footprint_path, 0) 
        driver_kml = ogr.GetDriverByName('KML')
        dataset_kml = driver_kml.CreateDataSource(self.kml_footprint_path)
        if dataset_shp is not None:
            for layer_index in range(dataset_shp.GetLayerCount()):
                layer = dataset_shp.GetLayerByIndex(layer_index)
                dataset_kml.CopyLayer(layer, layer.GetName())
            arcpy.AddMessage(f"Conversion in kml completed! File saved at: {self.kml_footprint_path}")
        else:
            arcpy.AddError("Kml file could not be produced")
        
class QualityCheckFootprintTool:

    def __init__(self):
        self.EMSR_Number = arcpy.GetParameterAsText(0)
        self.AOI = arcpy.GetParameterAsText(1)
        self.map_type = arcpy.GetParameterAsText(2)
        self.sensor_user_input = arcpy.GetParameterAsText(3)
        self.sensor_resolution = arcpy.GetParameter(10)
        self.sensor_folder = arcpy.GetParameterAsText(4)
        self.acquisition = arcpy.GetParameterAsText(7)
        self.availability = arcpy.GetParameterAsText(8)
        self.reception = arcpy.GetParameterAsText(9)
        self.pre_image_condition = arcpy.GetParameter(6)
        self.image_path = arcpy.GetParameterAsText(5)
        self.clean_A0 = arcpy.GetParameter(11)
        self.sensor_name_convention = {"COSMO-SkyMed": "CSK",
        "ESRI World Imagery": "ESRIWI",
        "GeoEye-1": "GEOEYE1",
        "Landsat-8": "LANDSAT8",
        "Landsat-9": "LANDSAT9",
        "PAZ": "PAZ",
        "Pleiades Neo": "PNEO",
        "PlanetScope": "PLANETSCOPE",
        "Pleiades": "PLEIADES",
        "RADARSAT-2": "RADARSAT2",
        "Sentinel-1": "SENTINEL1",
        "Sentinel-2": "SENTINEL2",
        "SkySat": "SKYSAT",
        "SPOT-6": "SPOT6",
        "SPOT-7": "SPOT7",
        "WorldView-2": "WORLDVIEW2",
        "WorldView-3": "WORLDVIEW3",
        "TerraSAR-X": "TERRASARX",
        "ALOS": "ALOS",
        "COSMO-SkyMed Second Generation": "CSG",
        "Drone Imagery":"DRONEIMAGERY",
        "GeoSat-1": "GEOSAT1",
        "GeoSat-2": "GEOSAT2",
        "ICEYE": "ICEYE",
        "IKONOS": "IKONOS1",
        "ISAT": "ISAT",
        "Kompsat-2": "KOMPSAT2",
        "Kompsat-3": "KOMPSAT3",
        "Kompsat-5": "KOMPSAT5",
        "Kompsat-8": "KOMPSAT8",
        "QuickBird": "QUICKBIRD",
        "RapidEye": "RAPIDEYE",
        "Resourcesat-1": "RESOURCESAT1",
        "Resourcesat-2": "RESOURCESAT2",
        "Sentinel-3": "SENTINEL3",
        "TIANHUI-1": "TIANHUI1",
        "WorldView-1": "WORLDVIEW2", 
        "Landsat-8": "LANDSAT8",
        "Landsat-9": "LANDSAT9",
        "Plane Imagery" : "PLANEIMAGERY"     
        }
        self.sensor_domain_value = {"COSMO-SkyMed": 89,
        "ESRI World Imagery": 17,
        "GeoEye-1": 25,
        "Landsat-8": 38,
        "Landsat-9": 38,
        "PAZ": 81,
        "Pleiades Neo": 83,
        "PlanetScope": 76,
        "Pleiades": 45,
        "RADARSAT-2": 50,
        "Sentinel-1": 55,
        "Sentinel-2": 56,
        "SkySat": 85,
        "SPOT-6": 62,
        "SPOT-7": 62,
        "WorldView-2": 74,
        "WorldView-3": 75,
        "TerraSAR-X": 66,
        "ALOS": 1,
        "COSMO-SkyMed Second Generation": 89,
        "Drone Imagery": 79,
        "GeoSat-1": 86,
        "GeoSat-2": 87,
        "ICEYE": 90,
        "IKONOS": 30,
        "ISAT": 88,
        "Kompsat-2": 33,
        "Kompsat-3": 34,
        "Kompsat-5": 35,
        "Kompsat-8": 997, # Must be added to domain values
        "QuickBird": 49,
        "RapidEye": 51,
        "Resourcesat-1": 52,
        "Resourcesat-2": 52,
        "Sentinel-3": 77,
        "TIANHUI-1": 67,
        "WorldView-1": 73,
        "Landsat-8": 38,
        "Landsat-9": 997, # Must be added to domain values
        "Plane Imagery" : 78
        }
        self.PrePost = "pre" if self.pre_image_condition else "post"
        self.SAR = (True if self.sensor_user_input in ["COSMO-SkyMed","PAZ","RADARSAT-2", "Sentinel-1", "TerraSAR-X"]
        else False)
        self.preOrthoFolder, self.postOrthoFolder, self.imageQCFolder, self.storing_folder= self.Create02IMAGESDirs()
        self.A0_source_table ,self.A2_image_footprint, self.A1_Area_of_Interest, self.gdb_path = self.RetreiveLayers() 
        self.activation_crs = cems_utils.getUTMZoneGpd(self.A1_Area_of_Interest)
        self.productionSite = self.ProductionSite() 
        self.sensor_name = self.sensor_name_convention[self.sensor_user_input]
        self.metadata_reader = MetadataReader(self.sensor_folder)
        self.DateTime = self.GetAcquisitionDateTime()
        self.filename_word, self.filename_footprint = self.GetFileName()
        self.shapefile_footprint_path = os.path.join(self.storing_folder, self.filename_footprint + '.shp')
        self.kml_footprint_path = os.path.join(self.imageQCFolder, self.filename_footprint + '.kml')
        self.QualityCheckCompiler = QualityCheckCompiler(self.EMSR_Number, self.AOI, self.map_type, self.filename_word, 
        self.DateTime, self.availability, self.reception, 
        self.sensor_resolution, self.imageQCFolder, self.productionSite)
        self.ImageFootprint = ImageFootprint(self.image_path, self.sensor_folder, self.preOrthoFolder, self.postOrthoFolder, self.shapefile_footprint_path, 
        self.kml_footprint_path,self.A2_image_footprint, 
        self.A0_source_table, self.gdb_path, self.pre_image_condition, self.sensor_user_input, self.DateTime, 
        self.sensor_resolution, self.clean_A0, self.filename_word, 
        self.sensor_domain_value, self.SAR, self.activation_crs) 
        
    def InitialCheck(self):
        attribute1 = "source_nam"
        attribute2 = "src_date"
        attribute3 = "source_tm"
        Date = self.DateTime[:10]
        Time = self.DateTime[11:]
        processed = False
        with arcpy.da.UpdateCursor(self.A0_source_table, ["OID@", attribute1, attribute2, attribute3]) as cursor:
            for row in cursor:
                if row[1] == self.sensor_domain_value[self.sensor_user_input] and row[2] == Date and row[3] == Time:
                    processed = True
                    arcpy.AddWarning("This image was already recorded in A0_source before you launched this tool. Make sure the automatically assigned or_src_id is correct!")
                    break
        word_document = os.path.join(self.imageQCFolder, self.filename_word + '_IMAGE_QC.docx')
        kml_footprint = os.path.join(self.imageQCFolder, self.filename_footprint +'.kml')
        shapefile_footprint = os.path.join(self.storing_folder, self.filename_footprint + '.shp')
        ORTHO_image = os.path.join(self.postOrthoFolder, self.filename_word + '.tif')
        files = [word_document, kml_footprint, ORTHO_image, shapefile_footprint]
        if processed:
            for file in files:
                if os.path.exists(file):
                    processed = True
                else:
                    processed = False
                    break
        return(processed)

    def Create02IMAGESDirs(self):
        aprx = arcpy.mp.ArcGISProject('current')
        original_path = aprx.filePath
        O3MAPS_path = os.path.dirname(original_path)
        EMSR_path = os.path.dirname(O3MAPS_path)
        output_path = EMSR_path + "\\02IMAGES"
        preRawFolder = reduce(os.path.join, [output_path, "PRE", "RAW"])
        preOrthoFolder = reduce(os.path.join, [output_path, "PRE", "ORTHO"]) 
        postRawFolder = reduce(os.path.join, [output_path, "POST", "RAW"])
        postOrthoFolder = reduce(os.path.join, [output_path, "POST", "ORTHO"]) 
        imageQCFolder = reduce(os.path.join, [output_path, "QC"])
        path_of_storing_folder = (f"AOI{self.AOI}_{self.map_type}_PRE_IMAGE" if self.pre_image_condition 
        else f"AOI{self.AOI}_{self.map_type}_POST_IMAGE")
        storing_folder = reduce(os.path.join, [output_path, "Tool_QC_outputs", path_of_storing_folder])
        folders = [preRawFolder, preOrthoFolder, postRawFolder, postOrthoFolder, imageQCFolder, storing_folder]
        arcpy.AddMessage("Checking if all necessary folders are present in your system...")
        for f in folders:
            if not os.path.exists(f):
                arcpy.AddMessage(f'{f} path not found, creating it...')
                os.makedirs(f)
        return(preOrthoFolder, postOrthoFolder, imageQCFolder, storing_folder)

    def RetreiveLayers(self):
        aprx = arcpy.mp.ArcGISProject("current")
        map_display = aprx.listMaps('Map Display')[0]
        A1_Area_of_Iterest = [lyr for lyr in map_display.listLayers() if lyr.name == "General Information"][0]
        A0_source_table = [table for table in map_display.listTables() if table.name == "A0_source_a"][0]
        A2_image_footprint = [lyr for lyr in map_display.listLayers() if lyr.name == "Image Footprint"][0]
        A2_image_footprint_path = A2_image_footprint.dataSource
        A0_source_table_path = A0_source_table.dataSource
        gdb_path = os.path.split(A2_image_footprint.dataSource)[0]
        return(A0_source_table_path, A2_image_footprint_path, A1_Area_of_Iterest , gdb_path)
    
    def ProductionSite(self):
        aprx = arcpy.mp.ArcGISProject('current')
        layout = aprx.listLayouts('EMS_Template_v*')[0]
        textLayoutElm = layout.listElements('TEXT_ELEMENT', 'publicationText')[0]
        productionSite = re.split('Map produced by | released by', textLayoutElm.text)[1]
        return(productionSite)

    def GetAcquisitionDateTime(self):
        arcpy.AddMessage(f"Reading {self.sensor_user_input} metadata...")
        if self.sensor_user_input == "Sentinel-2":
            DateTime = self.metadata_reader.Sentinel2()
        elif self.sensor_user_input == "Sentinel-1":
            DateTime = self.metadata_reader.Sentinel1()
        elif self.sensor_user_input in ["WorldView-2", "WorldView-3", "GeoEye-1"]:
            DateTime = self.metadata_reader.WorldView_and_Geoeye()
        elif self.sensor_user_input in ["SPOT-6", "SPOT-7", "Pleiades"]:
            DateTime = self.metadata_reader.SPOT_and_Pleiads()
        elif self.sensor_user_input == "Pleiades Neo":
            DateTime = self.metadata_reader.Pleiads_neo()
        elif self.sensor_user_input == "COSMO-SkyMed":
            DateTime = self.metadata_reader.CosmoSkyMed()
        elif self.sensor_user_input in ["PAZ", "PlanetScope", "SkySat", "TerraSAR-X"]:
            DateTime = self.metadata_reader.PAZ_and_Planetscope_and_SkySat_and_TERRASARX()
        elif self.sensor_user_input in ["Landsat-8", "Landsat-9"]:
            DateTime = self.metadata_reader.Landsat8_9()
        elif self.sensor_user_input == "RADARSAT-2":
            DateTime = self.metadata_reader.RadarSat2()
        else:
            DateTime = self.acquisition
        return(DateTime)
    
    def GetFileName(self):
        year = self.DateTime[6:10]
        month = self.DateTime[3:5]
        day = self.DateTime[0:2]
        hour = self.DateTime[11:13]
        min = self.DateTime[14:16]
        if self.map_type in ["REF", "FEP", "DEL", "GRA"]:
            File_name_word = f"{self.EMSR_Number}_AOI{self.AOI}_{self.map_type[0:3]}_PRODUCT_{self.sensor_name}_{year}{month}{day}_{hour}{min}_ORTHO"
            File_name_footprint = File_name_word.replace("ORTHO", "IMAGE_FOOTPRINT")
        else:
            File_name_word = f"{self.EMSR_Number}_AOI{self.AOI}_{self.map_type[0:3]}_{self.map_type[3:]}_{self.sensor_name}_{year}{month}{day}_{hour}{min}_ORTHO"
            File_name_footprint = File_name_word.replace("ORTHO", "IMAGE_FOOTPRINT")
        return(File_name_word, File_name_footprint)
    
    def RunTool(self):
        arcpy.AddMessage("Checking if this image has already been processed...")
        processed = self.InitialCheck()
        if processed == False:
            arcpy.AddMessage("This image hasn't been processed yet. Proceeding with operations...")
            composite_images = self.ImageFootprint.Composite()
            self.ImageFootprint.MosaickAndExport(composite_images)
            self.ImageFootprint.RemoveInvalidData()
            self.ImageFootprint.ShapefileFootprint()
            abort_procedure = self.ImageFootprint.UploadOnToc()
            if abort_procedure == False:
                self.ImageFootprint.KmlFootprint()
                or_src_id, tuples_list = self.ImageFootprint.A0SourceCompiler()
                self.ImageFootprint.A2_eraser(tuples_list)
                self.ImageFootprint.FindNoData(or_src_id)
                self.ImageFootprint.GDBFootprint()
                self.QualityCheckCompiler.Compiler()
                arcpy.AddMessage(f"REMEMBER: You can find a .shp footprint backup at: {self.storing_folder}")
            else:
                arcpy.AddWarning("Procedure aborted by the user!")
        else:
            arcpy.AddError("All the files relative to this image are already present in your system!")

if __name__ == "__main__":
    QualityCheckFootprintTool().RunTool()